"""
Mall AI Word Report: scraped data + optional web research → OpenAI → .docx export.
"""

import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

# Optional web research via DuckDuckGo (no API key)
try:
    from duckduckgo_search import DDGS
    HAS_DUCKDUCKGO = True
except ImportError:
    HAS_DUCKDUCKGO = False

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Reuse OpenAI from llm_engine
from llm_engine import _call_openai_chat


def _is_likely_tenant_name(name: str) -> bool:
    """True if name looks like a tenant/shop name, not a post caption or URL."""
    if not name or not isinstance(name, str):
        return False
    s = (name or "").strip()
    if len(s) < 2 or len(s) > 80:
        return False
    s_lower = s.lower()
    if any(x in s_lower for x in ("http", "instagram.com", "facebook.com", " | ", "reel", "sponsored")):
        return False
    return True


def _gather_web_research(
    mall_name: str,
    new_shop_names: list,
    vacated_shop_names: list,
    all_tenant_names_from_data: Optional[list] = None,
    max_results_per_query: int = 5,
    max_queries: int = 12,
) -> str:
    """
    Run web search queries using the data (mall name, new/vacated tenants, full tenant list)
    and return concatenated snippets for context. Uses DuckDuckGo (no API key).
    """
    if not HAS_DUCKDUCKGO:
        return ""

    snippets = []
    mall_clean = (mall_name or "shopping mall").strip()
    if not mall_clean:
        mall_clean = "shopping mall"

    queries = [
        f"new stores openings {mall_clean}",
        f"coming soon tenants {mall_clean}",
        f"latest news {mall_clean} mall",
        f"retail openings {mall_clean}",
    ]
    if new_shop_names:
        names = [n for n in new_shop_names[:3] if n and _is_likely_tenant_name(str(n))]
        if names:
            queries.append(f"{' '.join(names)} store opening {mall_clean}")
    if vacated_shop_names:
        names = [n for n in vacated_shop_names[:2] if n and _is_likely_tenant_name(str(n))]
        if names:
            queries.append(f"{mall_clean} store closure {names[0]}")
    # Use full tenant/shop list from data to search the web (not just scraped summary)
    if all_tenant_names_from_data:
        # Pick a sample of tenant names that look like real shops (not post text)
        sample = [n for n in all_tenant_names_from_data if _is_likely_tenant_name(str(n))][:10]
        for name in sample[:4]:  # up to 4 extra queries from data
            short = (name or "").strip()[:50]
            if short:
                queries.append(f"{short} {mall_clean} store")
                queries.append(f"{short} retail opening")

    seen = set()
    unique_queries = []
    for q in queries:
        q = q.strip()
        if q and q not in seen and len(unique_queries) < max_queries:
            seen.add(q)
            unique_queries.append(q)

    try:
        with DDGS() as ddgs:
            for q in unique_queries:
                try:
                    for r in ddgs.text(q, max_results=max_results_per_query):
                        title = (r.get("title") or "").strip()
                        body = (r.get("body") or "").strip()
                        href = (r.get("href") or "").strip()
                        if title or body:
                            snippets.append(f"[{title}]\n{body}\nSource: {href}")
                except Exception:
                    continue
    except Exception:
        return ""

    if not snippets:
        return ""
    return "\n\n---\n\n".join(snippets[:20])  # cap total snippets


def _build_context(
    scraped_df=None,
    structured_data=None,
    llm_json=None,
    input_url: str = "",
    web_research_text: str = "",
) -> str:
    """Build a single context string for the OpenAI prompt."""
    parts = []

    # 1) Scraped data summary
    if scraped_df is not None and not scraped_df.empty:
        cols = list(scraped_df.columns)
        if "shop_name" in cols:
            shops = scraped_df["shop_name"].dropna().astype(str).tolist()
            parts.append("Scraped tenant list (from mall website/social):\n" + "\n".join(f"- {s}" for s in shops[:200]))
        if "source" in cols:
            parts.append("\nData sources: " + ", ".join(scraped_df["source"].dropna().unique().tolist()))
    else:
        parts.append("(No scraped dataframe provided)")

    # 2) Structured comparison (new / vacated / shifted)
    if structured_data and isinstance(structured_data, dict):
        stats = structured_data.get("stats", {})
        parts.append("\n\nStructured comparison stats: " + json.dumps(stats, indent=2))
        new_shops = structured_data.get("new_shops", [])
        vacated = structured_data.get("vacated_shops", [])
        shifted = structured_data.get("shifted_shops", [])
        if new_shops:
            parts.append("\nNew shops: " + json.dumps(new_shops, indent=2))
        if vacated:
            parts.append("\nVacated shops: " + json.dumps(vacated, indent=2))
        if shifted:
            parts.append("\nShops that changed floor: " + json.dumps(shifted, indent=2))

    # 3) Existing LLM analysis (overall report)
    if llm_json and isinstance(llm_json, dict):
        parts.append("\n\nExisting AI analysis (overall): " + json.dumps(llm_json.get("overall", llm_json), indent=2))
        if llm_json.get("metadata"):
            parts.append("\nMetadata: " + json.dumps(llm_json["metadata"], indent=2))

    # 4) Input URLs
    if input_url:
        parts.append("\n\nInput URL(s) used for scraping: " + input_url.strip())

    # 5) Web research
    if web_research_text:
        parts.append("\n\n--- Web research (snippets from internet search) ---\n" + web_research_text)

    return "\n".join(parts)


def _call_openai_for_report(context: str, web_research_included: bool) -> str:
    """
    Ask OpenAI to generate a structured report (markdown with ## sections)
    from the provided data and web research snippets.
    """
    instruction = (
        "Using the data below (mall/tenant data and, when provided, web research snippets from internet search), "
        "write a professional mall research report. Use clear headings and bullet points. "
    )
    if web_research_included:
        instruction += (
            "The 'Web research' section contains real search results from the web: use it to add points of interest, "
            "recent news, store openings/closures, and context about tenants. Weave web research with the tenant data. "
        )
    instruction += (
        "Output format: use markdown with ## for main sections and ### for subsections. "
        "Use bullet points (- or *) for lists. Do not invent data; only use what is provided in the data and web research.\n\n"
    )

    prompt = instruction + "DATA:\n" + context + "\n\nRequired sections in your report:\n" + """
## Executive Summary
(2-4 sentences on occupancy trend and key changes)

## New Tenants – Points of Interest
(For each new shop: name, floor if known, and a short point of interest or context; if web research has info on a brand, include it here)

## Structured Changes
- New shops (list and count)
- Vacated / closed shops (list and count)
- Shops that changed floor (if any)

## Insights and Recommendations
(Business insights and any recommendations based on the data and, if provided, web research)

## Metadata
(Mall name, report date, data sources)
"""

    raw = _call_openai_chat(
        prompt,
        temperature=0.2,
        max_tokens=4096,
        response_format=None,  # plain text / markdown
        timeout_seconds=180,
    )
    return raw or ""


def _docx_from_markdown_report(report_text: str) -> BytesIO:
    """
    Parse markdown-style report (## headings, bullets) and build a Word document.
    """
    doc = Document()
    doc.add_heading("Mall AI Research Report", 0)

    # Normalize: \r\n -> \n
    text = (report_text or "").replace("\r\n", "\n").strip()
    if not text:
        doc.add_paragraph("No report content generated.")
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # Split into blocks by double newline, then detect headings vs paragraphs
    blocks = re.split(r"\n(?=\s*#+\s)", text)
    # First block might be before any ##
    if blocks and blocks[0].strip() and not re.match(r"^\s*#+\s", blocks[0]):
        for line in blocks[0].strip().split("\n"):
            line = line.strip()
            if line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:].strip(), style="List Bullet")
            elif line:
                doc.add_paragraph(line)
        blocks = blocks[1:]

    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        first = lines[0].strip()
        # ## Heading -> add_heading(..., level=1)  ### -> level=2
        if first.startswith("## "):
            doc.add_heading(first[3:].strip(), level=1)
            lines = lines[1:]
        elif first.startswith("### "):
            doc.add_heading(first[4:].strip(), level=2)
            lines = lines[1:]
        elif first.startswith("# "):
            doc.add_heading(first[2:].strip(), level=1)
            lines = lines[1:]

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if line_stripped.startswith("- ") or line_stripped.startswith("* "):
                doc.add_paragraph(line_stripped[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line_stripped)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_mall_word_report(
    scraped_df=None,
    structured_data=None,
    llm_json=None,
    input_url: str = "",
    do_web_research: bool = False,
    mall_name_for_search: Optional[str] = None,
    output_buffer: Optional[BytesIO] = None,
) -> BytesIO:
    """
    Generate a Word (.docx) report from scraped data, optional web research, and OpenAI.

    Args:
        scraped_df: DataFrame with scraped tenant data
        structured_data: Comparison result (new_shops, vacated_shops, etc.)
        llm_json: Existing LLM overall report (metadata + overall)
        input_url: URL(s) used for scraping
        do_web_research: If True, run web search and include snippets in context
        mall_name_for_search: Mall name to use for web search queries (default from llm_json or URL)
        output_buffer: Optional BytesIO to write to (otherwise creates new)

    Returns:
        BytesIO containing the .docx file.
    """
    if output_buffer is None:
        output_buffer = BytesIO()

    # Resolve mall name for web search
    mall_name = mall_name_for_search or ""
    if not mall_name and llm_json and isinstance(llm_json, dict):
        meta = llm_json.get("metadata", {})
        mall_name = (meta.get("mall_name") or "").strip()
    if not mall_name and input_url:
        try:
            from urllib.parse import urlparse
            host = urlparse(input_url.strip().split("\n")[0].strip()).netloc or ""
            mall_name = host.replace("www.", "").split(".")[0] or "mall"
        except Exception:
            mall_name = "mall"
    if not mall_name:
        mall_name = "shopping mall"

    # Web research: use data to drive search (mall + new/vacated + full tenant list from scraped data)
    web_research_text = ""
    if do_web_research:
        new_names = []
        vacated_names = []
        if structured_data and isinstance(structured_data, dict):
            new_names = [s.get("shop_name") for s in structured_data.get("new_shops", []) if s.get("shop_name")]
            vacated_names = [s.get("shop_name") for s in structured_data.get("vacated_shops", []) if s.get("shop_name")]
        all_tenant_names = []
        if scraped_df is not None and not scraped_df.empty and "shop_name" in scraped_df.columns:
            names = scraped_df["shop_name"].dropna().astype(str).str.strip().unique().tolist()
            all_tenant_names = [n for n in names if n and _is_likely_tenant_name(n)]
        web_research_text = _gather_web_research(
            mall_name, new_names, vacated_names, all_tenant_names_from_data=all_tenant_names
        )

    context = _build_context(
        scraped_df=scraped_df,
        structured_data=structured_data,
        llm_json=llm_json,
        input_url=input_url,
        web_research_text=web_research_text,
    )

    report_text = _call_openai_for_report(context, web_research_included=bool(web_research_text.strip()))
    if not report_text.strip():
        report_text = "Report could not be generated. Please check OpenAI API key and connection."

    docx_buffer = _docx_from_markdown_report(report_text)
    output_buffer.write(docx_buffer.getvalue())
    output_buffer.seek(0)
    return output_buffer
